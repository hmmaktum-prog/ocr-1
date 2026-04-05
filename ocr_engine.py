"""
OCR Engine — PaddleOCR-VL-1.5 integration with offline support.

Supported backends (in priority order):
  1. llama-cpp-server  — VL-1.5 GGUF via local llama.cpp server (Android / any device)
  2. PaddleOCRVL       — VL-1.5 via PaddlePaddle (desktop/server)
  3. Classic PaddleOCR — PP-OCRv4 Bengali (lightweight fallback)

Offline model cache:
  - GGUF models: ./models/gguf/
  - PaddlePaddle models: ./models/  (or $PADDLEOCR_MODEL_DIR)
"""

from __future__ import annotations  # BUG-26 fix: Python 3.8 compat for type hints

import os
import io
import re
import sys
import tempfile
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

MODEL_DIR = Path(os.environ.get("PADDLEOCR_MODEL_DIR", "./models"))
GGUF_DIR = MODEL_DIR / "gguf"

# Default llama-cpp-server URL (localhost on Android via Termux)
DEFAULT_LLAMA_SERVER_URL = "http://localhost:8111/v1"

_vl_pipeline = None
_classic_ocr = None


# ---------------------------------------------------------------------------
# Config (runtime-settable by the app)
# ---------------------------------------------------------------------------

class OCRConfig:
    """Runtime configuration for the OCR engine."""
    use_llama_server: bool = False
    llama_server_url: str = DEFAULT_LLAMA_SERVER_URL

_config = OCRConfig()


def set_llama_server(url: str):
    """Enable llama-cpp-server backend and set its URL."""
    _config.use_llama_server = True
    _config.llama_server_url = url
    global _vl_pipeline
    _vl_pipeline = None  # reset so it's re-initialised with new backend


def set_classic_mode():
    """Switch back to classic PP-OCRv4 Bengali."""
    _config.use_llama_server = False
    global _vl_pipeline
    _vl_pipeline = None


def get_active_mode() -> str:
    """Return a human-readable description of the current mode."""
    if _config.use_llama_server:
        return f"VL-1.5 (llama-cpp @ {_config.llama_server_url})"
    return "Classic PP-OCRv4 Bengali"


# ---------------------------------------------------------------------------
# Pipeline initialisation
# ---------------------------------------------------------------------------

def _init_vl_llama_pipeline():
    """VL-1.5 via llama-cpp-server (Android / offline)."""
    global _vl_pipeline
    if _vl_pipeline is not None:
        return _vl_pipeline
    try:
        from paddleocr import PaddleOCRVL
        logger.info("PaddleOCRVL: llama-cpp-server @ %s", _config.llama_server_url)
        _vl_pipeline = PaddleOCRVL(
            pipeline_version="v1.5",
            vl_rec_backend="llama-cpp-server",
            vl_rec_server_url=_config.llama_server_url,
        )
        logger.info("VL-1.5 (llama-cpp-server) pipeline ready.")
        return _vl_pipeline
    except Exception as exc:
        logger.warning("VL-1.5 llama-cpp-server init failed: %s", exc)
        return None


def _init_vl_paddle_pipeline():
    """VL-1.5 via PaddlePaddle (desktop)."""
    global _vl_pipeline
    if _vl_pipeline is not None:
        return _vl_pipeline
    try:
        from paddleocr import PaddleOCRVL

        vl_dir = MODEL_DIR / "PaddleOCR-VL-1.5"
        layout_dir = MODEL_DIR / "PP-DocLayoutV2"
        kwargs = {"pipeline_version": "v1.5"}
        if vl_dir.exists() and layout_dir.exists():
            logger.info("VL-1.5 Paddle: using local models in %s", MODEL_DIR)
            kwargs["vl_rec_model_dir"] = str(vl_dir)
            kwargs["layout_detection_model_dir"] = str(layout_dir)
        else:
            logger.info("VL-1.5 Paddle: downloading models on first run…")

        _vl_pipeline = PaddleOCRVL(**kwargs)
        logger.info("VL-1.5 (PaddlePaddle) pipeline ready.")
        return _vl_pipeline
    except Exception as exc:
        logger.warning("VL-1.5 PaddlePaddle init failed: %s", exc)
        return None


def _init_classic_ocr():
    """Classic PaddleOCR — PP-OCRv4 Bengali fallback."""
    global _classic_ocr
    if _classic_ocr is not None:
        return _classic_ocr
    try:
        from paddleocr import PaddleOCR

        det_dir = MODEL_DIR / "PP-OCRv4_mobile_det"
        rec_dir = MODEL_DIR / "ben_PP-OCRv4_rec"
        dict_file = MODEL_DIR / "ben_dict.txt"

        kwargs = {"use_angle_cls": True, "lang": "ben", "show_log": False}
        if det_dir.exists() and rec_dir.exists() and dict_file.exists():
            logger.info("Classic OCR: using local PP-OCRv4 models.")
            kwargs["det_model_dir"] = str(det_dir)
            kwargs["rec_model_dir"] = str(rec_dir)
            kwargs["rec_char_dict_path"] = str(dict_file)
        else:
            logger.info("Classic OCR: downloading PP-OCRv4 Bengali models…")

        _classic_ocr = PaddleOCR(**kwargs)
        logger.info("Classic PaddleOCR (PP-OCRv4, Bengali) ready.")
        return _classic_ocr
    except Exception as exc:
        logger.error("Classic OCR init failed: %s", exc)
        return None


def _get_vl_pipeline():
    """Return VL pipeline based on current config."""
    if _config.use_llama_server:
        return _init_vl_llama_pipeline()
    else:
        return _init_vl_paddle_pipeline()


# ---------------------------------------------------------------------------
# llama-cpp-server connectivity check
# ---------------------------------------------------------------------------

def test_llama_server(url: str | None = None) -> tuple:
    """
    Test connection to a llama-cpp-server.
    Returns (success: bool, message: str).
    """
    import urllib.request
    import urllib.error
    import json

    url = url or _config.llama_server_url
    base = url
    if base.endswith("/v1"):
        base = base[:-3]
    base = base.rstrip("/")
    health_url = base + "/health"

    try:
        with urllib.request.urlopen(health_url, timeout=5) as resp:
            data = json.loads(resp.read().decode())
            status = data.get("status", "unknown")
            if status in ("ok", "loading model", "no slot available"):
                return True, f"সার্ভার সংযুক্ত ({status})"
            return True, f"সার্ভার চলছে (status: {status})"
    except urllib.error.URLError as exc:
        return False, f"সংযোগ ব্যর্থ: {exc.reason}"
    except Exception as exc:
        return False, f"ত্রুটি: {exc}"


# ---------------------------------------------------------------------------
# Result extraction helpers
# ---------------------------------------------------------------------------

def _extract_markdown_from_vl_result(res) -> str:
    """Extract markdown text from a PaddleOCRVL result object."""
    markdown_text = ""

    # Attempt 1: json attribute
    try:
        data = getattr(res, "json", None)
        if isinstance(data, dict):
            markdown_text = (
                data.get("markdown")
                or data.get("text")
                or data.get("content")
                or ""
            )
            if not markdown_text:
                for key in ("parse_result", "result"):
                    sub = data.get(key)
                    if isinstance(sub, dict):
                        markdown_text = sub.get("markdown") or sub.get("text") or ""
                    elif isinstance(sub, str):
                        markdown_text = sub
                    if markdown_text:
                        break
        elif isinstance(data, str):
            markdown_text = data
    except Exception:
        pass

    # Attempt 2: save_to_markdown
    if not markdown_text:
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                res.save_to_markdown(save_path=tmpdir)
                for fname in sorted(os.listdir(tmpdir)):
                    if fname.endswith(".md"):
                        with open(os.path.join(tmpdir, fname), "r", encoding="utf-8") as f:
                            markdown_text = f.read()
                        break
        except Exception as exc:
            logger.debug("save_to_markdown failed: %s", exc)

    # Attempt 3: capture res.print() — BUG-20 fix: use StringIO without global redirect
    if not markdown_text:
        try:
            import contextlib
            buf = io.StringIO()
            with contextlib.redirect_stdout(buf):
                res.print()
            markdown_text = buf.getvalue()
        except Exception as exc:
            logger.debug("res.print() capture failed: %s", exc)

    return markdown_text.strip()


def _extract_text_from_classic_result(results) -> str:
    """Extract plain text from classic PaddleOCR result."""
    lines = []
    if not results:
        return ""
    for page_result in results:
        if not page_result:
            continue
        for word_info in page_result:
            if word_info and len(word_info) >= 2:
                text = (
                    word_info[1][0]
                    if isinstance(word_info[1], (list, tuple))
                    else str(word_info[1])
                )
                if text.strip():
                    lines.append(text.strip())
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Markdown → python-docx converter
# ---------------------------------------------------------------------------

def markdown_to_docx_page(doc, markdown_text: str, page_num: int):
    """
    Append a page's markdown content to a python-docx Document.
    Handles: headings, bold/italic, tables, lists, horizontal rules.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not markdown_text:
        doc.add_paragraph(f"[পৃষ্ঠা {page_num + 1}: কোনো টেক্সট পাওয়া যায়নি]")
        return

    lines = markdown_text.split("\n")
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        data_rows = [r for r in table_rows if not re.match(r"^\|[-| :]+\|$", r)]
        if not data_rows:
            in_table = False
            table_rows = []
            return
        parsed = []
        for row in data_rows:
            cells = [c.strip() for c in row.strip("|").split("|")]
            parsed.append(cells)
        if not parsed:
            in_table = False
            table_rows = []
            return
        col_count = max(len(r) for r in parsed)
        t = doc.add_table(rows=len(parsed), cols=col_count)
        t.style = "Table Grid"
        for r_idx, row in enumerate(parsed):
            for c_idx, cell_text in enumerate(row):
                if c_idx < col_count:
                    clean = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
                    clean = re.sub(r"\*(.+?)\*", r"\1", clean)
                    t.cell(r_idx, c_idx).text = clean
        in_table = False
        table_rows = []

    def strip_inline(text: str) -> str:
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
        text = re.sub(r"\[(.+?)\]\(.*?\)", r"\1", text)
        return text.strip()

    for line in lines:
        raw = line.rstrip()

        if raw.startswith("|"):
            if not in_table:
                in_table = True
            table_rows.append(raw)
            continue
        else:
            if in_table:
                flush_table()

        stripped = raw.strip()
        if not stripped:
            continue

        if re.match(r"^[-*_]{3,}$", stripped):
            doc.add_paragraph("─" * 40)
            continue

        if stripped.startswith("### "):
            doc.add_heading(strip_inline(stripped[4:]), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(strip_inline(stripped[3:]), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(strip_inline(stripped[2:]), level=1)
        elif re.match(r"^[-*+] ", stripped):
            doc.add_paragraph(strip_inline(stripped[2:]), style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            text = re.sub(r"^\d+\. ", "", stripped)
            doc.add_paragraph(strip_inline(text), style="List Number")
        else:
            text = strip_inline(stripped)
            if text:
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if in_table:
        flush_table()


# ---------------------------------------------------------------------------
# Direct llama.cpp HTTP OCR (Android-friendly, no PaddlePaddle needed)
# ---------------------------------------------------------------------------

def _ocr_image_via_llama_direct(img_bytes: bytes, server_url: str, task: str = "OCR") -> str:
    """
    Send a single image to llama.cpp OpenAI-compatible API and get OCR text.
    Uses only Python stdlib — works on Android without PaddlePaddle.

    task options: "OCR", "Table Recognition", "Formula Recognition"
    """
    import base64
    import json
    import urllib.request
    import urllib.error

    # Detect image type
    if img_bytes[:8] == b'\x89PNG\r\n\x1a\n':
        mime = "image/png"
    else:
        mime = "image/jpeg"

    b64_img = base64.b64encode(img_bytes).decode()

    payload = {
        "model": "PaddleOCR-VL-1.5",
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": f"{task}:"},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime};base64,{b64_img}"},
                    },
                ],
            }
        ],
        "max_tokens": 4096,
        "temperature": 0,
    }

    data = json.dumps(payload).encode("utf-8")
    api_url = server_url.rstrip("/") + "/chat/completions"
    req = urllib.request.Request(
        api_url,
        data=data,
        headers={"Content-Type": "application/json"},
    )

    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            result = json.loads(resp.read().decode("utf-8"))
            return result["choices"][0]["message"]["content"].strip()
    except urllib.error.URLError as exc:
        raise ConnectionError(
            f"llama.cpp সার্ভারে সংযোগ ব্যর্থ ({server_url}): {exc.reason}\n"
            "Termux-এ 'bash ~/start_vl_server.sh' চালান।"
        ) from exc


def _pdf_to_images(pdf_path: str, dpi: int = 150):
    """
    Render PDF pages to PIL Images using Android Native API, PyMuPDF, or pdf2image.
    Yields PIL Images one by one.
    """
    # 1. Try Android Native PdfRenderer via Pyjnius (No native C dependencies needed!)
    try:
        from jnius import autoclass
        from PIL import Image
        import io
        
        File = autoclass('java.io.File')
        ParcelFileDescriptor = autoclass('android.os.ParcelFileDescriptor')
        PdfRenderer = autoclass('android.graphics.pdf.PdfRenderer')
        Bitmap = autoclass('android.graphics.Bitmap')
        BitmapConfig = autoclass('android.graphics.Bitmap$Config')
        Color = autoclass('android.graphics.Color')
        CompressFormat = autoclass('android.graphics.Bitmap$CompressFormat')
        ByteArrayOutputStream = autoclass('java.io.ByteArrayOutputStream')

        file_obj = File(str(pdf_path))
        pfd = ParcelFileDescriptor.open(file_obj, ParcelFileDescriptor.MODE_READ_ONLY)
        renderer = PdfRenderer(pfd)
        
        page_count = renderer.getPageCount()
        
        for i in range(page_count):
            page = renderer.openPage(i)
            # Default PDF dpi is 72. Scale up
            scale = dpi / 72.0
            width = int(page.getWidth() * scale)
            height = int(page.getHeight() * scale)
            
            bitmap = Bitmap.createBitmap(width, height, BitmapConfig.ARGB_8888)
            bitmap.eraseColor(Color.WHITE)
            
            page.render(bitmap, None, None, PdfRenderer.Page.RENDER_MODE_FOR_DISPLAY)
            
            stream = ByteArrayOutputStream()
            bitmap.compress(CompressFormat.PNG, 100, stream)
            image_bytes = bytes(stream.toByteArray())
            
            img = Image.open(io.BytesIO(image_bytes))
            yield img
            
            page.close()
            bitmap.recycle()
            stream.close()
            
        renderer.close()
        pfd.close()
        return
    except ImportError:
        pass  # Not on Android
    except Exception as exc:
        logger.warning("Android Native PDF renderer failed: %s", exc)

    # 2. Try PyMuPDF (Desktop)
    try:
        import fitz
        from PIL import Image

        # BUG-22 fix: use context manager for safe document close
        with fitz.open(pdf_path) as pdf_doc:
            for page in pdf_doc:
                mat = fitz.Matrix(dpi / 72, dpi / 72)
                pix = page.get_pixmap(matrix=mat)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                yield img
        return
    except ImportError:
        pass
    except Exception:
        pass

    # Fallback: pdf2image (desktop only) — returns list, wrap as generator
    try:
        from pdf2image import convert_from_path
        for img in convert_from_path(pdf_path, dpi=dpi):
            yield img
        return
    except Exception as exc:
        raise RuntimeError(f"PDF পৃষ্ঠা পড়তে ব্যর্থ: {exc}") from exc


def _pdf_page_count(pdf_path: str) -> int:
    """Get total page count without rendering (for progress reporting)."""
    # 1. Android Native
    try:
        from jnius import autoclass
        File = autoclass('java.io.File')
        ParcelFileDescriptor = autoclass('android.os.ParcelFileDescriptor')
        PdfRenderer = autoclass('android.graphics.pdf.PdfRenderer')
        file_obj = File(str(pdf_path))
        pfd = ParcelFileDescriptor.open(file_obj, ParcelFileDescriptor.MODE_READ_ONLY)
        renderer = PdfRenderer(pfd)
        count = renderer.getPageCount()
        renderer.close()
        pfd.close()
        return count
    except Exception:
        pass

    # 2. PyMuPDF fallback
    try:
        import fitz
        with fitz.open(pdf_path) as doc:
            return len(doc)
    except Exception:
        pass
    try:
        from pdf2image import pdfinfo_from_path
        info = pdfinfo_from_path(pdf_path)
        return info.get("Pages", 0)
    except Exception:
        return 0


def _image_to_png_bytes(image) -> bytes:
    """Convert PIL Image to PNG bytes."""
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Public OCR API
# ---------------------------------------------------------------------------

def ocr_image_bytes(img_bytes: bytes) -> str:
    """
    Run OCR on a single image (bytes). Returns markdown/plain text.

    Priority:
      1. llama-cpp-server direct HTTP (Android-safe, no PaddlePaddle)
      2. PaddleOCRVL pipeline (desktop/server, needs PaddlePaddle)
      3. Classic PP-OCRv4 Bengali (lightweight fallback)
    """
    # 1. Direct llama.cpp HTTP (primary for Android)
    if _config.use_llama_server:
        try:
            return _ocr_image_via_llama_direct(img_bytes, _config.llama_server_url)
        except Exception as exc:
            logger.warning("Direct llama OCR failed: %s", exc)
            # BUG-17 fix: skip VL pipeline when in llama mode, go straight to classic
            classic = _init_classic_ocr()
            if classic is not None:
                try:
                    results = classic.ocr(img_bytes, cls=True)
                    return _extract_text_from_classic_result(results)
                except Exception as exc2:
                    logger.error("Classic OCR also failed: %s", exc2)
            return ""

    # 2. PaddleOCRVL pipeline (desktop — only when NOT in llama mode)
    vl = _get_vl_pipeline()
    if vl is not None:
        try:
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp.write(img_bytes)
                tmp_path = tmp.name
            try:
                output = list(vl.predict(tmp_path))
                if output:
                    return _extract_markdown_from_vl_result(output[0])
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass
        except Exception as exc:
            logger.warning("VL pipeline inference failed: %s — falling back.", exc)

    # 3. Classic OCR
    classic = _init_classic_ocr()
    if classic is not None:
        try:
            results = classic.ocr(img_bytes, cls=True)
            return _extract_text_from_classic_result(results)
        except Exception as exc:
            logger.error("Classic OCR also failed: %s", exc)

    return ""


def ocr_pdf(pdf_path: str, progress_callback=None) -> list:
    """
    Run OCR on a PDF file.
    Returns list of (page_num, markdown_text) tuples.

    Priority:
      1. Direct llama.cpp server — page by page (Android + desktop)
      2. PaddleOCRVL pipeline — full PDF (desktop)
      3. Classic PP-OCRv4 — page by page (desktop fallback)
    """
    pdf_path = str(pdf_path)
    results = []

    # --- 1. Direct llama.cpp HTTP (Android-safe) ---
    if _config.use_llama_server:
        try:
            logger.info("Direct llama.cpp OCR: %s", pdf_path)
            total = _pdf_page_count(pdf_path) or 1
            for page_num, image in enumerate(_pdf_to_images(pdf_path)):
                if progress_callback:
                    progress_callback(page_num, total)
                img_bytes = _image_to_png_bytes(image)
                del image  # free memory immediately
                try:
                    md_text = _ocr_image_via_llama_direct(
                        img_bytes, _config.llama_server_url
                    )
                except Exception as exc:
                    logger.warning("Page %d llama OCR failed: %s", page_num, exc)
                    md_text = f"[পৃষ্ঠা {page_num + 1}: সার্ভার ত্রুটি — {exc}]"
                del img_bytes  # free memory immediately
                results.append((page_num, md_text))
            return results
        except Exception as exc:
            logger.warning("Direct llama PDF OCR failed: %s — falling back.", exc)
            results = []

    # --- 2. PaddleOCRVL pipeline (desktop, full PDF) ---
    if not _config.use_llama_server:  # BUG-17 fix: only try VL when not in llama mode
        vl = _get_vl_pipeline()
        if vl is not None:
            try:
                logger.info("VL pipeline: processing %s", pdf_path)
                output = list(vl.predict(pdf_path))
                total = len(output)
                for page_num, res in enumerate(output):
                    if progress_callback:
                        progress_callback(page_num, total)
                    md_text = _extract_markdown_from_vl_result(res)
                    results.append((page_num, md_text))
                if results:
                    return results
            except Exception as exc:
                logger.warning("VL pipeline PDF processing failed: %s — falling back.", exc)
                results = []

    # --- 3. Classic PaddleOCR (desktop fallback) ---
    logger.info("Fallback: classic PaddleOCR page-by-page.")
    total = _pdf_page_count(pdf_path) or 1

    try:
        image_gen = _pdf_to_images(pdf_path)
    except Exception as exc:
        logger.error("Cannot read PDF pages: %s", exc)
        return [(0, f"[PDF পড়তে ব্যর্থ: {exc}]")]

    classic = _init_classic_ocr()

    for page_num, image in enumerate(image_gen):
        if progress_callback:
            progress_callback(page_num, total)

        img_bytes = _image_to_png_bytes(image)
        del image  # free memory immediately
        text = ""
        if classic is not None:
            try:
                res = classic.ocr(img_bytes, cls=True)
                text = _extract_text_from_classic_result(res)
            except Exception as exc:
                logger.warning("Classic OCR page %d failed: %s", page_num, exc)

        del img_bytes  # free memory immediately
        results.append((page_num, text or f"[পৃষ্ঠা {page_num + 1}: কোনো টেক্সট পাওয়া যায়নি]"))

    return results


def build_docx_from_ocr_results(ocr_results: list):
    """Convert list of (page_num, markdown_text) into a python-docx Document."""
    from docx import Document

    doc = Document()
    total = len(ocr_results)
    for page_num, markdown_text in ocr_results:
        doc.add_heading(f"পৃষ্ঠা {page_num + 1}", level=2)
        markdown_to_docx_page(doc, markdown_text, page_num)
        if page_num < total - 1:
            doc.add_page_break()
    return doc
